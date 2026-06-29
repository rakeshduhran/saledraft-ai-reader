import base64
import io
import os
from typing import List, Tuple

import fitz  # PyMuPDF
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from PIL import Image

APP_TITLE = "SaleDraft AI Hindi Text Reader"
MODEL = os.getenv("AI_MODEL", "gpt-4o-mini")
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "8"))
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY")) if os.getenv("OPENAI_API_KEY") else None

app = FastAPI(title=APP_TITLE)
app.mount("/static", StaticFiles(directory="static"), name="static")

SYSTEM_PROMPT = """
You are an expert Hindi-English document text reader for Indian registry, jamabandi, khasra, gazette and legal documents.
Task: Extract ONLY the visible text. Preserve Hindi as Hindi/Devanagari and English as English. Do not translate.
Keep line breaks and page order. Correct obvious OCR spelling only when context is certain. Do not invent missing text.
For unclear words, write [अस्पष्ट]. Keep numbers, slashes, khasra numbers and punctuation exactly as seen.
""".strip()


def html_page() -> str:
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()


@app.get("/", response_class=HTMLResponse)
def home():
    return html_page()


def image_to_data_url(img: Image.Image) -> str:
    # limit large images for API cost/speed while keeping readable quality
    max_side = 1800
    img = img.convert("RGB")
    w, h = img.size
    if max(w, h) > max_side:
        scale = max_side / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=88)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return f"data:image/jpeg;base64,{b64}"


def ai_read_images(images: List[Image.Image], mode: str = "registry") -> str:
    if not client:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY server पर set नहीं है।")
    content = [{"type": "input_text", "text": SYSTEM_PROMPT + f"\nMode: {mode}\nReturn clean extracted text only."}]
    for img in images:
        content.append({"type": "input_image", "image_url": image_to_data_url(img)})
    resp = client.responses.create(model=MODEL, input=[{"role": "user", "content": content}])
    return resp.output_text.strip()


def direct_pdf_text(data: bytes, page_range: str = "") -> Tuple[str, int]:
    doc = fitz.open(stream=data, filetype="pdf")
    pages = parse_pages(page_range, len(doc))
    text_parts = []
    for pno in pages:
        page = doc[pno]
        txt = page.get_text("text") or ""
        if txt.strip():
            text_parts.append(f"--- Page {pno+1} ---\n{txt.strip()}")
    return "\n\n".join(text_parts).strip(), len(pages)


def parse_pages(page_range: str, total: int) -> List[int]:
    if not page_range.strip():
        return list(range(min(total, MAX_PDF_PAGES)))
    pages = []
    for part in page_range.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            for n in range(int(a), int(b) + 1):
                if 1 <= n <= total:
                    pages.append(n - 1)
        else:
            n = int(part)
            if 1 <= n <= total:
                pages.append(n - 1)
    return pages[:MAX_PDF_PAGES]


def pdf_pages_to_images(data: bytes, page_range: str = "") -> List[Image.Image]:
    doc = fitz.open(stream=data, filetype="pdf")
    pages = parse_pages(page_range, len(doc))
    images = []
    for pno in pages:
        page = doc[pno]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images


@app.post("/extract-image")
async def extract_image(file: UploadFile = File(...), mode: str = Form("registry")):
    data = await file.read()
    try:
        img = Image.open(io.BytesIO(data))
    except Exception:
        raise HTTPException(status_code=400, detail="Image file read नहीं हो पाई।")
    text = ai_read_images([img], mode=mode)
    return {"text": text, "method": "AI Vision Image Reader"}


@app.post("/extract-pdf")
async def extract_pdf(file: UploadFile = File(...), pages: str = Form(""), mode: str = Form("registry"), direct_first: str = Form("true")):
    data = await file.read()
    if direct_first == "true":
        txt, count = direct_pdf_text(data, pages)
        # if useful text found, return it without AI cost
        if len(txt) > 40:
            return {"text": txt, "method": f"Direct PDF Text ({count} pages)"}
    images = pdf_pages_to_images(data, pages)
    if not images:
        raise HTTPException(status_code=400, detail="PDF pages नहीं मिलीं।")
    text = ai_read_images(images, mode=mode)
    return {"text": text, "method": "AI Vision PDF Scan Reader"}


@app.post("/cleanup")
async def cleanup_text(text: str = Form(...), mode: str = Form("registry")):
    if not client:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY server पर set नहीं है।")
    prompt = f"""Hindi-English mixed registry document text को साफ करें।
Rules:
- Hindi को Hindi ही रखें, English को English ही रखें, translation नहीं।
- khasra numbers, dates, slashes, page labels exact रखें।
- केवल obvious OCR mistakes सुधारें।
- कोई नया तथ्य न जोड़ें।
Mode: {mode}

TEXT:\n{text}"""
    resp = client.responses.create(model=MODEL, input=prompt)
    return {"text": resp.output_text.strip(), "method": "AI Hindi Cleanup"}


@app.post("/download-docx")
async def download_docx(text: str = Form(...)):
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    for para in text.split("\n"):
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=extracted_text.docx"})


@app.post("/download-txt")
async def download_txt(text: str = Form(...)):
    data = text.encode("utf-8")
    return StreamingResponse(io.BytesIO(data), media_type="text/plain; charset=utf-8", headers={"Content-Disposition": "attachment; filename=extracted_text.txt"})
